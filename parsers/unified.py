"""
Parsers unificados: qualquer formato de entrada (pdf, csv, docx, xlsx, md, txt)
é convertido para uma estrutura intermediária comum:

{
    "raw_header_lines": [...],   # linhas de texto livre antes da tabela (fornecedor, cnpj, cidade, pedido/cotacao)
    "rows": [ {"codigo":..., "nome":..., "un":..., "qtd":..., "valor":...}, ... ],
    "source_file": "nome.ext"
}

Nada de matching ou classificação acontece aqui — só extração bruta.
"""
import csv
import io
import os
import re

import openpyxl
import docx as docx_lib

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False


COLUMN_ALIASES = {
    "codigo": {"codigo", "código", "cod"},
    "nome": {"produto", "descricao", "descrição"},
    "un": {"un", "unidade", "und"},
    "qtd": {"qtd", "qtde", "quantidade"},
    "valor": {"valor", "preco", "preço", "val liq", "val unit", "valor unitario", "valor unitário"},
}


def _normalize_col(col):
    col = (col or "").strip().lower()
    for canon, aliases in COLUMN_ALIASES.items():
        if col in aliases:
            return canon
    return None


def _table_from_grid(grid):
    """grid: list of list of cell strings. Detects header row and maps columns."""
    header_idx = None
    col_map = {}
    for i, row in enumerate(grid):
        mapped = {j: _normalize_col(str(c)) for j, c in enumerate(row)}
        hits = sum(1 for v in mapped.values() if v)
        if hits >= 3:  # heuristic: at least 3 recognized columns = header row
            header_idx = i
            col_map = {v: j for j, v in mapped.items() if v}
            break
    if header_idx is None:
        return [], []

    header_lines = []
    for row in grid[:header_idx]:
        text = " ".join(str(c) for c in row if c).strip()
        if text:
            header_lines.append(text)

    rows = []
    for row in grid[header_idx + 1:]:
        if not any(str(c).strip() for c in row if c is not None):
            continue
        entry = {}
        for canon, j in col_map.items():
            if j < len(row):
                entry[canon] = row[j]
        if entry.get("nome"):
            rows.append(entry)
    return header_lines, rows


def parse_csv(path):
    with open(path, newline="", encoding="utf-8") as f:
        sample = f.read(2048)
        f.seek(0)
        delimiter = ";" if sample.count(";") > sample.count(",") else ","
        text = f.read()

    # header metadata lines are before the actual delimited table
    lines = text.splitlines()
    table_start = None
    for i, line in enumerate(lines):
        if delimiter in line and any(_normalize_col(c) for c in line.split(delimiter)):
            table_start = i
            break
    header_lines = [l.strip() for l in lines[:table_start] if l.strip()] if table_start else []
    table_text = "\n".join(lines[table_start:]) if table_start is not None else text
    reader = csv.reader(io.StringIO(table_text), delimiter=delimiter)
    grid = list(reader)
    _, rows = _table_from_grid(grid)
    return {"raw_header_lines": header_lines, "rows": rows, "source_file": os.path.basename(path)}


def parse_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    grid = [list(row) for row in ws.iter_rows(values_only=True)]
    grid = [row for row in grid if any(c is not None for c in row)]
    grid = [[("" if c is None else c) for c in row] for row in grid]
    header_lines, rows = _table_from_grid(grid)
    return {"raw_header_lines": header_lines, "rows": rows, "source_file": os.path.basename(path)}


def parse_docx(path):
    d = docx_lib.Document(path)
    header_lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    rows = []
    for table in d.tables:
        grid = [[c.text for c in row.cells] for row in table.rows]
        _, table_rows = _table_from_grid(grid)
        rows.extend(table_rows)
    return {"raw_header_lines": header_lines, "rows": rows, "source_file": os.path.basename(path)}


def parse_md_or_txt(path):
    with open(path, encoding="utf-8") as f:
        text = f.read()
    lines = text.splitlines()

    # markdown pipe table?
    pipe_lines = [i for i, l in enumerate(lines) if "|" in l]
    if pipe_lines:
        table_start = pipe_lines[0]
        header_lines = [l.strip() for l in lines[:table_start] if l.strip()]
        grid = []
        for l in lines[table_start:]:
            if not l.strip():
                continue
            if re.match(r"^\s*\|?[\s:|-]+\|?\s*$", l):
                continue  # markdown separator row like |---|---|
            cells = [c.strip() for c in l.strip().strip("|").split("|")]
            grid.append(cells)
        _, rows = _table_from_grid(grid)
        return {"raw_header_lines": header_lines, "rows": rows, "source_file": os.path.basename(path)}

    # tab-separated?
    tab_lines = [i for i, l in enumerate(lines) if "\t" in l]
    if tab_lines:
        table_start = tab_lines[0]
        header_lines = [l.strip() for l in lines[:table_start] if l.strip()]
        grid = [l.split("\t") for l in lines[table_start:] if l.strip()]
        _, rows = _table_from_grid(grid)
        return {"raw_header_lines": header_lines, "rows": rows, "source_file": os.path.basename(path)}

    return {"raw_header_lines": [l.strip() for l in lines if l.strip()], "rows": [], "source_file": os.path.basename(path)}


def parse_pdf(path):
    """Handles the 'Ordem de Venda' style documents: free text + tabular item lines."""
    text = ""
    if HAS_PDFPLUMBER:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                text += t + "\n"
    else:
        raise RuntimeError("pdfplumber not available")

    lines = [l for l in text.splitlines() if l.strip()]
    header_lines = []
    rows = []
    # Ordem de venda item lines: pdfplumber puts the product name on the line(s)
    # immediately BEFORE the numeric row, e.g.:
    #   "FERRO CA50 10 mm 3/8 KG"
    #   "1 05509 753,00 KG R$ 16,00 R$ 8,19 R$ 7,81 F R$ 5.880,93"
    item_re = re.compile(
        r"^\s*(\d+)\s+(\S+)\s+([\d.,]+)\s+([A-Z][A-Z0-9]{1,2})\s+(?:(.*?)\s+)?R\$\s*([\d.,]+)\s+R\$\s*([\d.,]+)\s+R\$\s*([\d.,]+)\s+[A-Z]\s+R\$\s*([\d.,]+)\s*$"
    )
    name_buffer = []
    STOP_PREFIXES = ("Vendedor", "FORMA", "ENTREGA", "OBSERVA", "Total", "Usuário",
                      "ITEM CÓDIGO", "ORDEM DE VENDA", "NÃO É", "SITUAÇÃO", "Loja:", "End.:",
                      "Cliente", "Endereço", "Cidade", "CNPJ/CPF", "liberação")

    for line in lines:
        m = item_re.match(line)
        if m:
            item_no, codigo, qtd, un, nome_extra, val_unit, val_desc, val_liq, val_total = m.groups()
            parts = name_buffer + ([nome_extra.strip()] if nome_extra else [])
            nome = " ".join(p for p in parts if p).strip() or "(produto não identificado)"
            rows.append({
                "codigo": codigo,
                "nome": nome,
                "un": un,
                "qtd": qtd,
                "valor": val_liq,
                "valor_total": val_total,
            })
            name_buffer = []
        elif line.strip().startswith(STOP_PREFIXES):
            header_lines.append(line)
            name_buffer = []
        else:
            name_buffer.append(line.strip())

    return {"raw_header_lines": header_lines, "rows": rows, "source_file": os.path.basename(path)}


def parse_any(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        return parse_csv(path)
    if ext == ".xlsx":
        return parse_xlsx(path)
    if ext == ".docx":
        return parse_docx(path)
    if ext in (".md", ".txt"):
        return parse_md_or_txt(path)
    if ext == ".pdf":
        return parse_pdf(path)
    raise ValueError(f"Formato não suportado: {ext}")
